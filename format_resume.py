from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def rgb_color_from_hex(hex_color):
    hex_color = hex_color.lstrip('#')
    return RGBColor(*(int(hex_color[i:i+2], 16) for i in (0, 2, 4)))

def set_strict_page_margins_fixed(section, inches=1):
    section.top_margin = Inches(inches)
    section.bottom_margin = Inches(inches)
    section.left_margin = Inches(inches)
    section.right_margin = Inches(inches)

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_horizontal_line(paragraph, color="auto"):
    p = paragraph._p  # Access the internal lxml element
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)

def add_header_with_fully_flush_left_logo(doc, logo_path, bar_color_hex):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    section = doc.sections[0]
    header = section.header
    section.header_distance = Inches(0.15)
    for para in header.paragraphs:
        p = para._element
        p.getparent().remove(p)
    logo_para = header.add_paragraph()
    logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    logo_para.paragraph_format.left_indent = Inches(-0.7)
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after = Pt(0)
    run = logo_para.add_run()
    run.add_picture(logo_path, width=Inches(1))
    bar_para = header.add_paragraph()
    bar_para.paragraph_format.space_before = Pt(0)
    bar_para.paragraph_format.space_after = Pt(0)
    bar_para.paragraph_format.left_indent = Inches(-0.7)
    bar_para.paragraph_format.right_indent = Inches(-0.7)
    bar_para.paragraph_format.line_spacing = Pt(1)
    bar_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    border_xml = (
        f'<w:pBorders {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="40" w:color="{bar_color_hex.lstrip("#")}" w:space="0"/>'
        f'</w:pBorders>'
    )
    bar_para._p.get_or_add_pPr().append(parse_xml(border_xml))
    return doc

def create_resume_doc(name, summary, certifications, skills, experience, education, logo_path):
    hex_teal = "#284b62"
    hex_dark = "#0b233b"
    doc = Document()
    set_strict_page_margins_fixed(doc.sections[0], inches=1)

    # Add header with logo and bar
    add_header_with_fully_flush_left_logo(doc, logo_path, hex_dark)

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(name.upper())
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = rgb_color_from_hex(hex_dark)

    # Summary
    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(12)
    summary_para.add_run(summary)

    # Certifications and Skills
    certs_skills_para = doc.add_paragraph()
    certs_skills_para.paragraph_format.space_after = Pt(12)
    certs_skills_para_run = certs_skills_para.add_run(" | ".join(certifications + skills))
    certs_skills_para_run.font.name = 'Calibri'
    certs_skills_para_run.font.size = Pt(10)
    certs_skills_para_run.font.bold = True
    certs_skills_para_run.font.color.rgb = rgb_color_from_hex(hex_dark)

    # Experience Header
    exp_header = doc.add_paragraph()
    exp_header_run = exp_header.add_run("EXPERIENCE")
    exp_header_run.font.name = 'Calibri'
    exp_header_run.font.size = Pt(19)
    exp_header_run.font.bold = True
    exp_header_run.font.color.rgb = rgb_color_from_hex(hex_dark)
    insert_horizontal_line(exp_header, hex_dark)

    # Experience entries
    for job in experience:
        # Company, City, State & Date on same line with tab stops
        comp_para = doc.add_paragraph()
        comp_para.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        comp_para.paragraph_format.space_before = Pt(2)
        comp_para.paragraph_format.space_after = Pt(0)
        comp_run = comp_para.add_run(
            ", ".join(filter(None, [job.get('company', ''), job.get('city', ''), job.get('state', '')]))
        )
        comp_run.font.name = 'Calibri'
        comp_run.font.size = Pt(11)
        comp_run.font.bold = True
        comp_run.font.color.rgb = rgb_color_from_hex(hex_teal)
        # Tab to date
        comp_para.add_run("\t")
        date_run = comp_para.add_run(job.get('years', ''))
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(11)
        date_run.font.bold = True
        date_run.font.color.rgb = rgb_color_from_hex(hex_teal)
        # Job title (italic and colored)
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(3)
        title_run = title_para.add_run(job.get('title', ''))
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(11)
        title_run.font.italic = True
        title_run.font.color.rgb = rgb_color_from_hex(hex_teal)
        # Bullet points
        for bullet in job.get('bullets', []):
            bullet_para = doc.add_paragraph(style='ListBullet')
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            bullet_para.paragraph_format.space_before = Pt(2)
            bullet_para.paragraph_format.space_after = Pt(2)
            bullet_run = bullet_para.add_run(bullet)
            bullet_run.font.name = 'Calibri'
            bullet_run.font.size = Pt(10)

    # Education Header
    edu_header = doc.add_paragraph()
    edu_header_run = edu_header.add_run("EDUCATION")
    edu_header_run.font.name = 'Calibri'
    edu_header_run.font.size = Pt(14)
    edu_header_run.font.bold = True
    edu_header_run.font.color.rgb = rgb_color_from_hex(hex_dark)
    insert_horizontal_line(edu_header, hex_dark)

    # Education entries
    for edu in education:
        univ_para = doc.add_paragraph()
        univ_run = univ_para.add_run(edu.get('university', ''))
        univ_run.font.name = 'Calibri'
        univ_run.font.size = Pt(11)
        univ_run.font.bold = True
        univ_run.font.color.rgb = rgb_color_from_hex(hex_teal)
        deg_para = doc.add_paragraph()
        deg_para.paragraph_format.space_before = Pt(0)
        deg_run = deg_para.add_run(edu.get('degree', ''))
        deg_run.italic = True
        deg_run.font.name = 'Calibri'
        deg_run.font.size = Pt(11)
        deg_run.font.color.rgb = rgb_color_from_hex(hex_teal)

    return doc
